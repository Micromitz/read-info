# conver-docx-to-csv.py
# Uso:
#   CSV:  python conver-docx-to-csv.py "C:/ruta/archivo.docx" --csv "C:/ruta/salida.csv"
#   JSON: python conver-docx-to-csv.py "C:/ruta/archivo.docx" --json "C:/ruta/salida.json"


# Esta versión extrae "id, tipo método, URL endpoint, Código EH, fecha ejecución, y cantidad vulnerabilidades (críticas, altas, medias y bajas)" en CSV y JSON

import re
import csv
import json
import argparse
import unicodedata
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Tuple
from docx import Document

HTTP_METHODS = {"GET","POST","PUT","PATCH","DELETE","OPTIONS","HEAD"}
SPANISH_MONTHS = {
    "enero":1,"febrero":2,"marzo":3,"abril":4,"mayo":5,"junio":6,
    "julio":7,"agosto":8,"septiembre":9,"setiembre":9,"octubre":10,
    "noviembre":11,"diciembre":12
}

# ---------------- Utilitarios ----------------
def _norm(s: str) -> str:
    if s is None: return ""
    s = s.replace("\xa0"," ").strip()
    s = "".join(c for c in unicodedata.normalize("NFKD", s) if not unicodedata.combining(c))
    return s.lower()

def _clean(s: str) -> str:
    return " ".join((s or "").replace("\xa0"," ").split())

def parse_spanish_date(text: str) -> str:
    """'03 de diciembre de 2024' -> '03/12/2024' (maneja también 02-dic-2024)."""
    if not text: return ""
    t = _norm(text)
    m = re.search(r'(\d{1,2})\s+de\s+([a-z]+)\s+de\s+(\d{4})', t)
    if not m:
        m = re.search(r'(\d{1,2})[-/ ]([a-z]{3,})[-/ ](\d{4})', t)
        if not m: return ""
    day, mon_txt, year = m.group(1), m.group(2), m.group(3)
    mon_txt = mon_txt.strip(".").lower()
    month = None
    for k in list(SPANISH_MONTHS.keys()):
        if mon_txt.startswith(k[:3]):
            month = SPANISH_MONTHS[k]; break
    if not month: month = SPANISH_MONTHS.get(mon_txt)
    if not month: return ""
    try:
        return datetime(int(year), int(month), int(day)).strftime("%d/%m/%Y")
    except Exception:
        return ""

# ---------------- Extractores ----------------
def extract_eh_code(doc: Document) -> str:
    pat = re.compile(r'\bEH-[A-Z0-9]{2,}-\d{5}\b', re.IGNORECASE)
    for p in doc.paragraphs:
        m = pat.search(p.text)
        if m: return m.group(0).upper()
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                m = pat.search(c.text)
                if m: return m.group(0).upper()
    return ""

def extract_fecha_version(doc: Document) -> str:
    for t in doc.tables:
        for r in t.rows:
            if len(r.cells) < 2: continue
            if "fecha de la version" in _norm(r.cells[0].text):
                return parse_spanish_date(r.cells[1].text)
    for p in doc.paragraphs:
        if "fecha" in _norm(p.text):
            parsed = parse_spanish_date(p.text)
            if parsed: return parsed
    return ""

def _is_resume_table(t) -> bool:
    if t.rows:
        header = " | ".join(_norm(c.text) for c in t.rows[0].cells)
        return "severidad" in header and ("cvss" in header or "vulnerabilidad" in header)
    return False

def count_vulnerabilities_by_severity(doc: Document) -> Dict[str,int]:
    counts = {"CRITICA":0, "ALTA":0, "MEDIA":0, "BAJA":0}
    for t in doc.tables:
        if not _is_resume_table(t): continue
        for row in t.rows[1:]:
            if not row.cells: continue
            sev = _clean(row.cells[-1].text).upper().replace("Í","I")
            for key in counts.keys():
                if key in sev:
                    counts[key] += 1
                    break
    return counts

def extract_endpoints(doc: Document) -> List[Tuple[str,str,str]]:
    LABELS = {"punto de entrada":"num","metodo":"metodo","método":"metodo","uri":"uri"}
    rows = []
    for t in doc.tables:
        data = {"num":"","metodo":"","uri":""}
        for r in t.rows:
            if len(r.cells) < 2: continue
            key = _norm(r.cells[0].text)
            if key in LABELS:
                data[LABELS[key]] = _clean(r.cells[1].text)
        if data["num"] and data["metodo"] and data["uri"]:
            metodo = data["metodo"].split()[0].upper()
            if metodo in HTTP_METHODS:
                rows.append((data["num"], metodo, data["uri"]))
    def sort_key(r):
        try: return int(re.sub(r"\D","", r[0]))
        except: return r[0]
    rows.sort(key=sort_key)
    return rows

# ---------------- Main ----------------
def main():
    ap = argparse.ArgumentParser(description="Convierte .docx a CSV/JSON con EH, fecha y conteos de severidad.")
    ap.add_argument("docx", help="Ruta del .docx de entrada")
    ap.add_argument("--csv", help="Ruta del CSV de salida")
    ap.add_argument("--json", help="Ruta del JSON de salida")
    args = ap.parse_args()

    doc = Document(args.docx)
    eh = extract_eh_code(doc)
    fecha = extract_fecha_version(doc)
    counts = count_vulnerabilities_by_severity(doc)
    endpoints = extract_endpoints(doc)

    if args.csv:
        with open(args.csv, "w", newline="", encoding="utf-8") as f:
            w = csv.writer(f, delimiter=";")
            for num, metodo, uri in endpoints:
                w.writerow([num, metodo, uri, eh, fecha, counts["CRITICA"], counts["ALTA"], counts["MEDIA"], counts["BAJA"]])
        print(f"CSV generado: {args.csv}")

    if args.json:
        data = []
        for num, metodo, uri in endpoints:
            data.append({
                "num": num,
                "metodo": metodo,
                "uri": uri,
                "eh_code": eh,
                "fecha_ejec": fecha,
                "vulnerabilidades": {
                    "criticas": counts["CRITICA"],
                    "altas": counts["ALTA"],
                    "medias": counts["MEDIA"],
                    "bajas": counts["BAJA"]
                }
            })
        with open(args.json, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        print(f"JSON generado: {args.json}")

if __name__ == "__main__":
    main()
